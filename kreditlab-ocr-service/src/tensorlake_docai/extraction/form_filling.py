# SPDX-License-Identifier: Apache-2.0
from __future__ import annotations

import asyncio
import json
import math
import os
from dataclasses import dataclass, asdict
from typing import Dict


from PIL import Image
import base64
import tempfile
from tensorlake.applications import cls, function
from tensorlake_docai.vlm.workflow_images import file_convertion_image
from tensorlake_docai.pipeline.api import Usage, ParsedDocumentRef
from tensorlake_docai.models.intermediate_objects import ParseResult, FormFillingResult
from tensorlake_docai.pipeline.output_formatter import format_final_output

# --- Data Models ---


@dataclass
class BoundingBox:
    x1: float
    y1: float
    x2: float
    y2: float

    @property
    def center(self) -> tuple[float, float]:
        return ((self.x1 + self.x2) / 2, (self.y1 + self.y2) / 2)

    @property
    def width(self) -> float:
        return abs(self.x2 - self.x1)

    @property
    def height(self) -> float:
        return abs(self.y2 - self.y1)

    def distance_to(self, other: BoundingBox) -> float:
        cx1, cy1 = self.center
        cx2, cy2 = other.center
        return math.sqrt((cx1 - cx2) ** 2 + (cy1 - cy2) ** 2)

    def iou(self, other: BoundingBox) -> float:
        """Calculates Intersection over Union (IoU) with another box."""
        x_left = max(self.x1, other.x1)
        y_top = max(self.y1, other.y1)
        x_right = min(self.x2, other.x2)
        y_bottom = min(self.y2, other.y2)

        if x_right < x_left or y_bottom < y_top:
            return 0.0

        intersection_area = (x_right - x_left) * (y_bottom - y_top)

        area1 = self.width * self.height
        area2 = other.width * other.height

        union_area = area1 + area2 - intersection_area
        if union_area == 0:
            return 0.0

        return intersection_area / union_area

    def is_contained_in(self, other: BoundingBox, threshold: float = 0.9) -> bool:
        """Checks if this box is mostly contained within another box."""
        x_left = max(self.x1, other.x1)
        y_top = max(self.y1, other.y1)
        x_right = min(self.x2, other.x2)
        y_bottom = min(self.y2, other.y2)

        if x_right < x_left or y_bottom < y_top:
            return False

        intersection_area = (x_right - x_left) * (y_bottom - y_top)
        self_area = self.width * self.height
        if self_area == 0:
            return False
        return (intersection_area / self_area) >= threshold


@dataclass
class PageFragment:
    """Represents a text element or figure from the PDF parsing stage."""

    fragment_type: str
    content: str
    bbox: BoundingBox
    reading_order: int


@dataclass
class DetectedWidget:
    """Represents a form widget found by the object detector."""

    label: str  # e.g., 'text_input', 'checkbox'
    score: float
    bbox: BoundingBox
    linked_text: str | None = None  # The predicted label/question for this widget
    is_filled: bool = False  # Indicates if the detector thinks this is already filled
    page_number: int = 1
    # Indicates if the widget was already present in the source PDF
    is_existing: bool = False
    description: str | None = None
    text_content: str | None = None
    field_name: str | None = None
    # Context for disambiguation
    surrounding_text: str | None = None


@dataclass
class PageData:
    page_number: int
    fragments: list[PageFragment]
    # In a real scenario, this would hold the image reference for the detector
    image_path: str | None = None


# --- Core Logic Components ---


class LabelAssociator:
    """
    Responsible for linking detected widgets to their semantic text labels
    based on geometric proximity and alignment.
    """

    def associate(
        self,
        widget: DetectedWidget,
        fragments: list[PageFragment],
        is_docx: bool = False,
    ) -> str | None:
        """
        Finds the text fragment most likely associated with the widget.
        Prioritizes text to the Left (key-value pairs) or Above (headers).
        """
        best_candidate = None
        min_score = float("inf")

        w_x1, w_y1, w_x2, w_y2 = widget.bbox.x1, widget.bbox.y1, widget.bbox.x2, widget.bbox.y2
        w_center_y = (w_y1 + w_y2) / 2

        # Parameters based on document type to separate DOCX logic from PDF logic
        if is_docx:
            y_min_below = -50
            y_max_above = 300
            top_penalty = 50
        else:
            y_min_below = -20
            y_max_above = 100
            top_penalty = 500

        for frag in fragments:
            if frag.fragment_type not in ["text", "title", "section_header"]:
                continue

            f_x1, f_y1, f_x2, f_y2 = frag.bbox.x1, frag.bbox.y1, frag.bbox.x2, frag.bbox.y2
            f_center_y = (f_y1 + f_y2) / 2

            # Vertical alignment check
            y_diff = w_center_y - f_center_y

            if y_diff < y_min_below:
                continue  # Text is significantly below
            if y_diff > y_max_above:
                continue  # Text is significantly above

            # Horizontal relationship
            left_gap = w_x1 - f_x2  # Positive if text is to left
            right_gap = f_x1 - w_x2  # Positive if text is to right

            score = float("inf")

            # 1. Text is to the Left (Label: [Input])
            if left_gap >= -10:
                dist_score = left_gap if left_gap >= 0 else abs(left_gap) * 2
                v_penalty = abs(y_diff) * 2
                score = dist_score + v_penalty

            # 2. Text is Above (Label\n[Input])
            elif f_y2 <= w_y1 + 10:
                w_center_x = (w_x1 + w_x2) / 2
                f_center_x = (f_x1 + f_x2) / 2
                h_misalignment = abs(w_center_x - f_center_x)
                vertical_gap = w_y1 - f_y2
                score = vertical_gap + h_misalignment + top_penalty  # Prefer Left over Top

            # 3. Text is to the Right ([ ] Label)
            elif right_gap >= -10:
                dist_score = right_gap if right_gap >= 0 else abs(right_gap) * 2
                v_penalty = abs(y_diff) * 2
                if widget.label in ["checkbox", "radio"]:
                    score = dist_score + v_penalty + 50
                else:
                    score = (
                        dist_score + v_penalty + 1000
                    )  # Strong penalty for right-side text on inputs

            if score < min_score:
                min_score = score
                best_candidate = frag

        return best_candidate.content if best_candidate else None


class PdfFormAugmenter:
    """
    Handles the modification of the PDF to add interactive fields
    and generates debug information.
    """

    def augment_pdf(
        self,
        source_pdf_path: str,
        output_pdf_path: str,
        widgets: list[DetectedWidget],
    ):
        """
        Creates a new PDF with the additional widgets using pypdf.
        """
        import pypdf

        try:
            from pypdf.generic import (
                NameObject,
                DictionaryObject,
                ArrayObject,
                FloatObject,
                TextStringObject,
                NumberObject,
                BooleanObject,
            )

            reader = pypdf.PdfReader(source_pdf_path)
            writer = pypdf.PdfWriter()
            # Use clone_document_from_reader to properly copy the entire document structure,
            # including the AcroForm and its resources, which append_pages_from_reader does not.
            writer.clone_document_from_reader(reader)

            # Ensure we have an AcroForm dictionary to register fields
            if "/AcroForm" not in writer.root_object:
                writer.root_object[NameObject("/AcroForm")] = DictionaryObject()

            acroform = writer.root_object["/AcroForm"]

            # NeedAppearances is crucial for seeing the widgets without clicking them
            acroform[NameObject("/NeedAppearances")] = BooleanObject(True)

            # Ensure Default Resources (DR) are present for fonts used in DA strings
            if "/DR" not in acroform:
                acroform[NameObject("/DR")] = DictionaryObject()

            # Ensure Default Appearance (DA) is present globally
            if "/DA" not in acroform:
                acroform[NameObject("/DA")] = TextStringObject("/Helv 0 Tf 0 g")

            dr = acroform["/DR"]
            if "/Font" not in dr:
                dr[NameObject("/Font")] = DictionaryObject()

            font_dict = dr["/Font"]
            if "/Helv" not in font_dict:
                helv_font = DictionaryObject()
                helv_font[NameObject("/Type")] = NameObject("/Font")
                helv_font[NameObject("/Subtype")] = NameObject("/Type1")
                helv_font[NameObject("/BaseFont")] = NameObject("/Helvetica")
                font_dict[NameObject("/Helv")] = helv_font

            if "/Fields" not in acroform:
                acroform[NameObject("/Fields")] = ArrayObject()

            fields = acroform["/Fields"]

            for w in widgets:
                page_idx = w.page_number - 1
                if page_idx < 0 or page_idx >= len(writer.pages):
                    continue

                page = writer.pages[page_idx]
                mb = page.mediabox

                if w.is_existing:
                    if (w.is_filled or w.text_content) and "/Annots" in page:
                        page_top = float(mb.top)
                        page_left = float(mb.left)

                        for annot in page["/Annots"]:
                            obj = annot.get_object()
                            if obj.get("/Subtype") == "/Widget":
                                rect = obj.get("/Rect")
                                if not rect:
                                    continue
                                x_ll, y_ll, x_ur, y_ur = [float(c) for c in rect]
                                b_x1 = x_ll - page_left
                                b_y1 = page_top - y_ur

                                if abs(b_x1 - w.bbox.x1) < 1.0 and abs(b_y1 - w.bbox.y1) < 1.0:
                                    if w.label in ["checkbox", "radio"]:
                                        val = (
                                            NameObject("/Yes")
                                            if w.is_filled
                                            else NameObject("/Off")
                                        )
                                        obj[NameObject("/V")] = val
                                        obj[NameObject("/AS")] = val
                                    elif w.text_content:
                                        obj[NameObject("/V")] = TextStringObject(w.text_content)
                                        # Remove Appearance dictionary to force regeneration by the viewer
                                        if "/AP" in obj:
                                            del obj["/AP"]
                                    break
                    continue

                page_left = float(mb.left)
                page_bottom = float(mb.bottom)
                pg_w = float(mb.width)
                pg_h = float(mb.height)

                # Get rotation safely handling inheritance
                rotation = 0
                if hasattr(page, "rotation"):
                    rotation = page.rotation
                else:
                    rotation = page.get("/Rotate", 0)
                rotation = int(rotation) % 360

                # Transform coordinates from Image Space (Top-Left) to PDF Space (Bottom-Left, Unrotated)
                x1, y1, x2, y2 = w.bbox.x1, w.bbox.y1, w.bbox.x2, w.bbox.y2

                """
                if w.label == "text_input":
                    # Heuristic: Shrink the box slightly to prevent obscuring adjacent labels
                    # especially the label above the field.
                    w_box = abs(x2 - x1)
                    h_box = abs(y2 - y1)
                    x1 += w_box * 0.01
                    x2 -= w_box * 0.01
                    y1 += h_box * 0.10  # Shrink top by 10% (Image space y increases downwards)
                    y2 -= h_box * 0.02
                """
                bx_min, bx_max = min(x1, x2), max(x1, x2)
                by_min, by_max = min(y1, y2), max(y1, y2)

                # Ensure minimum dimensions for visibility
                if (bx_max - bx_min) < 10:
                    bx_max = bx_min + 10
                if (by_max - by_min) < 10:
                    by_max = by_min + 10

                if rotation == 0:
                    rect = [bx_min, pg_h - by_max, bx_max, pg_h - by_min]
                elif rotation == 90:
                    # Visual (0,0) -> PDF (0, 0)
                    rect = [by_min, bx_min, by_max, bx_max]
                elif rotation == 180:
                    # Visual (0,0) -> PDF (W, 0)
                    rect = [pg_w - bx_max, by_min, pg_w - bx_min, by_max]
                elif rotation == 270:
                    # Visual (0,0) -> PDF (W, H)
                    rect = [pg_w - by_max, pg_h - bx_max, pg_w - by_min, pg_h - bx_min]
                else:
                    rect = [bx_min, pg_h - by_max, bx_max, pg_h - by_min]

                # Apply offset
                rect = [
                    page_left + rect[0],
                    page_bottom + rect[1],
                    page_left + rect[2],
                    page_bottom + rect[3],
                ]

                # Create Widget

                # Create Annotation
                annot = DictionaryObject()
                annot[NameObject("/Type")] = NameObject("/Annot")
                annot[NameObject("/Subtype")] = NameObject("/Widget")
                annot[NameObject("/Rect")] = ArrayObject([FloatObject(c) for c in rect])

                annot[NameObject("/T")] = TextStringObject(w.field_name)
                annot[NameObject("/F")] = NumberObject(4)  # Print flag

                # Border Style
                bs = DictionaryObject()
                bs[NameObject("/W")] = NumberObject(1)
                bs[NameObject("/S")] = NameObject("/S")
                annot[NameObject("/BS")] = bs

                # Appearance Characteristics (Black Border)
                mk = DictionaryObject()
                mk[NameObject("/BC")] = ArrayObject(
                    [FloatObject(0), FloatObject(0), FloatObject(0)],
                )
                mk[NameObject("/BG")] = ArrayObject(
                    [FloatObject(0.9), FloatObject(0.95), FloatObject(1.0)],
                )

                mk[NameObject("/R")] = NumberObject(rotation)

                # Set MK dictionary
                annot[NameObject("/MK")] = mk

                if w.label in ["checkbox", "radio"]:
                    annot[NameObject("/FT")] = NameObject("/Btn")
                    if w.label == "radio":
                        annot[NameObject("/Ff")] = NumberObject(32768)
                    else:
                        annot[NameObject("/Ff")] = NumberObject(0)

                    if w.is_filled:
                        annot[NameObject("/V")] = NameObject("/Yes")
                        annot[NameObject("/AS")] = NameObject("/Yes")
                    else:
                        annot[NameObject("/V")] = NameObject("/Off")
                        annot[NameObject("/AS")] = NameObject("/Off")
                elif w.label == "signature":
                    annot[NameObject("/FT")] = NameObject("/Sig")
                else:
                    annot[NameObject("/FT")] = NameObject("/Tx")
                    annot[NameObject("/Ff")] = NumberObject(4096)  # Multiline
                    # Use 0 for font size to enable auto-sizing to fit the box
                    annot[NameObject("/DA")] = TextStringObject("/Helv 0 Tf 0 g")
                    if w.text_content:
                        annot[NameObject("/V")] = TextStringObject(w.text_content)

                # Link to the page to prevent "hanging" in some readers
                if hasattr(page, "indirect_ref"):
                    annot[NameObject("/P")] = page.indirect_ref
                elif hasattr(page, "indirectRef"):
                    annot[NameObject("/P")] = page.indirectRef

                # Use add_annotation to register the object and add to page
                annot_ref = writer.add_annotation(page_idx, annot)
                if annot_ref is None:
                    annot_ref = page["/Annots"][-1]

                # Also add to the global Fields array so it functions as a form
                fields.append(annot_ref)

            with open(output_pdf_path, "wb") as f:
                writer.write(f)
            print(f"Augmented PDF saved to {output_pdf_path}")

        except Exception as e:
            print(f"Error augmenting PDF: {e}")
            import traceback

            traceback.print_exc()


class DocxFormAugmenter:
    """
    Handles modification of DOCX files to add native form widgets.
    Uses text anchoring since DOCX is flow-based.
    """

    def _get_tc_text(self, tc) -> str:
        """Extracts all text from a <w:tc> element."""
        from docx.oxml.ns import qn

        text = ""
        if tc is None:
            return ""

        for p in tc.findall(qn("w:p")):
            for t in p.findall(f'.//{qn("w:t")}'):
                if t.text:
                    text += t.text
        return text

    def _set_tc_text(self, tc, text_to_add: str):
        """Clears and sets the text of a <w:tc> element."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        if tc is None:
            return
        # Clear existing content
        for p in tc.findall(qn("w:p")):
            tc.remove(p)
        # Create new content
        new_p = OxmlElement("w:p")
        new_r = OxmlElement("w:r")
        new_t = OxmlElement("w:t")
        new_t.text = text_to_add.strip()
        new_r.append(new_t)
        new_p.append(new_r)
        tc.append(new_p)

    def _build_visual_tc_grid(self, table) -> list[list[any]]:
        from docx.oxml.ns import qn

        try:
            num_cols = len(table._tbl.tblGrid.gridCol_lst)
        except AttributeError:
            num_cols = len(table.columns) if table.rows else 0

        num_rows = len(table.rows)
        grid = [[None for _ in range(num_cols)] for _ in range(num_rows)]

        for r, row in enumerate(table.rows):
            row_tc_elements = []
            for child in row._tr:
                if child.tag.endswith("}tc"):
                    row_tc_elements.append(child)
                elif child.tag.endswith("}sdt"):
                    sdt_content = child.find(qn("w:sdtContent"))
                    if sdt_content is not None:
                        tc_in_sdt = sdt_content.find(qn("w:tc"))
                        if tc_in_sdt is not None:
                            row_tc_elements.append(tc_in_sdt)

            grid_c = 0
            for tc in row_tc_elements:
                while grid_c < num_cols and grid[r][grid_c] is not None:
                    grid_c += 1

                if grid_c >= num_cols:
                    break

                tc_pr = tc.find(qn("w:tcPr"))

                # --- vMerge detection ---
                v_merge_elem = tc_pr.find(qn("w:vMerge")) if tc_pr is not None else None
                if v_merge_elem is not None:
                    val = v_merge_elem.get(qn("w:val"))  # "restart" or None
                    if val != "restart":
                        # Continuation cell — point to the cell above
                        if r > 0 and grid[r - 1][grid_c] is not None:
                            grid[r][grid_c] = grid[r - 1][grid_c]
                        grid_c += 1
                        continue

                # --- gridSpan detection ---
                grid_span = 1
                grid_span_elem = tc_pr.find(qn("w:gridSpan")) if tc_pr is not None else None
                if grid_span_elem is not None:
                    val_str = grid_span_elem.get(qn("w:val"))
                    try:
                        if val_str:
                            grid_span = int(val_str)
                    except (ValueError, TypeError):
                        pass

                for j in range(grid_span):
                    if (grid_c + j) < num_cols:
                        grid[r][grid_c + j] = tc

                grid_c += grid_span

        return grid

    def augment_docx(
        self,
        source_docx_bytes: bytes,
        output_docx_path: str,
        widgets: list[DetectedWidget],
        total_pdf_pages: int | None = None,
    ) -> list[str]:
        import tempfile
        import os
        import re
        import subprocess
        from docx import Document
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        temp_docx_path = None
        warnings = []

        def normalize_text(text: str | None) -> str:
            if not text:
                return ""
            text = text.lower().strip()
            text = re.sub(r"[●·]", "•", text)  # Normalize bullets
            text = re.sub(r"[\u2010-\u2015\u2212\uFE58\uFE63\uFF0D]", "-", text)
            text = re.sub(r"\s*:\s*", ":", text)
            # Keep alphanumeric, spaces, colon, hyphen, brackets, and normalized bullet
            text = re.sub(r"[^\w\s:\-\[\]\(\)•]", "", text)
            return " ".join(text.split())

        def map_paragraphs_to_list_numbers(doc):
            """Pre-calculate list numbers for all paragraphs to avoid O(N^2) complexity."""
            para_index_to_num = {}
            counters = {}
            for i, para in enumerate(doc.paragraphs):
                pPr = para._element.find(qn("w:pPr"))
                if pPr is None:
                    continue
                numPr = pPr.find(qn("w:numPr"))
                if numPr is None:
                    continue

                ilvl_el = numPr.find(qn("w:ilvl"))
                numId_el = numPr.find(qn("w:numId"))
                if ilvl_el is None or numId_el is None:
                    continue

                ilvl = int(ilvl_el.get(qn("w:val")))
                numId = int(numId_el.get(qn("w:val")))

                key = (numId, ilvl)
                counters[key] = counters.get(key, 0) + 1

                para_index_to_num[i] = counters[key]

            return para_index_to_num

        def has_numPr(paragraph):
            """
            Check if paragraph has numbering properties in XML.
            """
            # If we are just replacing a placeholder, we might modify the p element.
            # But checking numPr helps identify if it WAS a list item.
            # However, in python-docx, if we access p._element, we are fine.
            if paragraph is None:
                return False

            pPr = paragraph._element.find(qn("w:pPr"))
            if pPr is None:
                return False
            return pPr.find(qn("w:numPr")) is not None

        def check_context_similarity(
            paragraph_text: str, context_text: str | None, current_xml_num: int | None = None
        ) -> tuple[float, list[str]]:
            """
            Calculates similarity between paragraph text and widget's surrounding text context.
            Returns a score between 0.0 and 1.0.
            """
            details = []
            if not context_text:
                return 0.5, ["No context"]

            p_norm = normalize_text(paragraph_text)
            c_norm = normalize_text(context_text)

            if not c_norm:
                return 0.5, ["Empty normalized context"]

            # Jaccard similarity on words
            p_words = set(p_norm.split())
            c_words = set(c_norm.split())

            if not c_words:
                return 0.5, ["No words in context"]

            intersection = p_words.intersection(c_words)
            base_score = len(intersection) / len(c_words)
            score = base_score
            details.append(f"Base: {base_score:.2f} ({len(intersection)}/{len(c_words)})")

            # Add a bonus for matching list numbers present in both contexts
            # Use original, un-normalized text for this check

            # Strict list number extraction: markers at start or delimited by punctuation (e.g. "1.", "(1)")
            c_nums = re.findall(r"(?:^|\s)[\(\[]?(\d+)[\)\]\.]", context_text)
            # Also check for simple number at start of context
            start_num = re.match(r"^\s*(\d+)\b", context_text)

            target_nums = set()
            if c_nums:
                target_nums.update(int(n) for n in c_nums)
            if start_num:
                target_nums.add(int(start_num.group(1)))

            if target_nums:
                details.append(f"Target nums: {list(target_nums)}")
                bonus = False
                # 1. Text-based number match
                p_num_match = re.match(r"^\s*[\(\[]?(\d+)[\)\]\.]?", paragraph_text.strip())
                if p_num_match and int(p_num_match.group(1)) in target_nums:
                    bonus = True
                    details.append(f"Num match (text) {p_num_match.group(1)}")

                # 2. XML-based number match (auto-numbering)
                if not bonus and current_xml_num is not None:
                    if current_xml_num in target_nums:
                        bonus = True
                        details.append(f"Num match (xml) {current_xml_num}")

                if bonus:
                    # Conditional bonus: huge bonus if text matches well, smaller if not
                    # This prevents "1. Unrelated text" from stealing "1. Related text"
                    if base_score > 0.3:
                        score += 1.5
                        details.append("Bonus +1.5 (High Sim)")
                    elif base_score > 0.1:
                        score += 0.5
                        details.append("Bonus +0.5 (Low Sim)")
                    else:
                        details.append("Bonus Skipped (Mismatch)")

            return score, details

        def map_paragraphs_to_pages(doc):
            # Keyed by paragraph XML element (_p) so that table-cell paragraphs are
            # covered correctly. doc.paragraphs returns all paragraphs in document
            # order (including inside tables), which the old body-child loop missed.
            _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            _br_page_xpath = f'.//{{{_W}}}br[@{{{_W}}}type="page"]'

            para_element_to_page = {}
            current_page = 1

            for para in doc.paragraphs:
                if para.paragraph_format.page_break_before:
                    current_page += 1

                para_element_to_page[para._p] = current_page

                # Count explicit manual page breaks inside this paragraph.
                # Use lxml findall instead of string-matching .xml to avoid false
                # positives and avoid serializing the full XML tree per paragraph.
                page_break_count = len(para._p.findall(_br_page_xpath))
                if page_break_count:
                    current_page += page_break_count

            return para_element_to_page

        # Static placeholder patterns shared across all process_paragraphs calls.
        # Checkbox-specific patterns are appended per-widget inside process_paragraphs.
        _BASE_PLACEHOLDER_PATTERNS = [
            re.compile(r"\[\s*[●•·]\s*\]"),
            re.compile(r"\[\s*_{3,}\s*\]"),
            re.compile(r"_{3,}"),
            re.compile(r"\.{3,}"),
            re.compile(r"-{3,}"),
            re.compile(r"\[\s*enter\s+text\s+here\s*\]", re.IGNORECASE),
            re.compile(r"click\s+here\s+to\s+enter\s+text", re.IGNORECASE),
            re.compile(r"click\s+here\s+to\s+enter\s+a\s+date", re.IGNORECASE),
            re.compile(r"\[[^\]]{1,80}\]"),  # broad bracket fallback
        ]
        _CHECKBOX_PLACEHOLDER_PATTERNS = [
            re.compile(r"\[\s*\]"),
            re.compile(r"\[x\]", re.IGNORECASE),
            re.compile(r"\u2610"),
            re.compile(r"\u2611"),
        ]

        def process_paragraphs(
            paragraphs,
            widget,
            para_page_map,
            max_doc_page,
            para_num_map,
            page_scale_factor=1.0,
            enable_page_bonus=True,
        ):
            # -------------------------------------------------------------------------
            # 1. Build value to insert
            # -------------------------------------------------------------------------
            if widget.label in ["checkbox", "radio"]:
                value_to_add = " \u2611" if widget.is_filled else " \u2610"
            elif widget.label == "signature" and widget.text_content:
                value_to_add = f" {widget.text_content}"
            elif widget.text_content:
                value_to_add = f" {widget.text_content.strip()}"
            else:
                value_to_add = ""

            def select_best_match(full_text, matches, context_text):
                if not matches:
                    return None
                # If only one match, return it
                if len(matches) == 1:
                    return matches[0]

                # Normalize context for keyword matching
                context_words = set(normalize_text(context_text).split())
                if not context_words:
                    return matches[0]

                best_match = matches[0]
                best_score = -1

                for m in matches:
                    # Look at text preceding the placeholder (up to 150 chars)
                    # This helps anchor the placeholder to its label (e.g. "Address: [•]")
                    start = max(0, m.start() - 150)
                    pre_text = full_text[start : m.start()]
                    pre_words = set(normalize_text(pre_text).split())

                    score = len(context_words.intersection(pre_words))
                    if score > best_score:
                        best_score = score
                        best_match = m
                return best_match

            # Combine surrounding text and linked label for richer context matching
            combined_context = (widget.surrounding_text or "") + " " + (widget.linked_text or "")
            combined_context_norm = normalize_text(combined_context)

            print(
                f"Processing widget '{widget.linked_text}' with label '{widget.label}' and value to add: '{value_to_add.strip()}'"
            )
            print(f"  > Context for matching: '{combined_context[:100]}...'")

            # -------------------------------------------------------------------------
            # 2. Placeholder patterns (ordered: most specific first)
            # -------------------------------------------------------------------------
            if widget.label in ["checkbox", "radio"]:
                placeholder_patterns = _BASE_PLACEHOLDER_PATTERNS + _CHECKBOX_PLACEHOLDER_PATTERNS
            else:
                placeholder_patterns = _BASE_PLACEHOLDER_PATTERNS

            # -------------------------------------------------------------------------
            # 3. Helper: replace text in paragraph preserving run formatting
            # -------------------------------------------------------------------------
            def apply_replacement(p, new_text):
                if p.runs:
                    p.runs[0].text = new_text
                    for run in p.runs[1:]:
                        run.text = ""
                else:
                    # Paragraph has no runs (e.g. inline SDT/content-control) —
                    # inject a bare run so the value is not silently dropped.
                    new_r = OxmlElement("w:r")
                    new_t = OxmlElement("w:t")
                    new_t.text = new_text
                    # Preserve leading/trailing whitespace that Word would otherwise strip.
                    new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                    new_r.append(new_t)
                    p._p.append(new_r)

            # -------------------------------------------------------------------------
            # 4. Case A — linked_text is itself a bracket placeholder e.g. "[in Italy]"
            #    The whole linked_text IS the thing to replace, no separate anchor needed
            # -------------------------------------------------------------------------
            linked = widget.linked_text.strip() if widget.linked_text else ""
            if re.match(r"^\[.+\]$", linked):
                candidates = []

                # Build flexible regex for bracketed placeholders
                inner_content = linked[1:-1].strip()
                is_generic_bullet = inner_content in ["•", "●", "·"]

                if is_generic_bullet:
                    # Special case for bullet placeholders: match any common bullet/placeholder style
                    escaped_linked = r"\[\s*[•●·\*\-_\u2022]\s*\]"
                else:
                    # Build a whitespace-flexible pattern for the bracket content.
                    # re.escape() doesn't escape spaces in Python 3.7+, so the old
                    # .replace(r'\ ', r'\s+') was a no-op. Join escaped words on \s+
                    # instead so internal whitespace variations always match.
                    inner_content = linked[1:-1]
                    escaped_inner = r"\s+".join(re.escape(w) for w in inner_content.split())
                    escaped_linked = r"\[\s*" + escaped_inner + r"\s*\]"

                direct_pattern = re.compile(escaped_linked, re.IGNORECASE)

                for i, p in enumerate(paragraphs):
                    if direct_pattern.search(p.text):
                        xml_num = para_num_map.get(i)
                        score, details = check_context_similarity(p.text, combined_context, xml_num)

                        # Page alignment bonus
                        p_page_num = para_page_map.get(p._p, 1)
                        if enable_page_bonus and max_doc_page > 1:
                            scaled_p_page = p_page_num * page_scale_factor
                            page_diff = abs(scaled_p_page - widget.page_number)

                            if page_diff <= 0.5:  # Close match (approx same page)
                                score += 2.0
                            elif page_diff <= 1.5:  # Adjacent page
                                score += 0.5

                        candidates.append((score, p, p_page_num, details))

                # Sort by score descending
                candidates.sort(key=lambda x: x[0], reverse=True)

                # Debug candidates
                if candidates and candidates[0][0] < 3.0:  # Only if score is not perfect
                    print(f"    > Top candidates for '{linked}':")
                    for score, p, p_page, dets in candidates[:3]:
                        print(
                            f"      - Score {score:.2f} (Page {p_page}): {p.text[:60].strip()}... {dets}"
                        )

                if candidates:
                    best_score, p, p_page, _ = candidates[0]
                    # If we have context, require a decent match, otherwise fallback to first found
                    if widget.surrounding_text and best_score < 0.1 and len(candidates) > 1:
                        print(
                            f"  > ⚠️ DOCX: Low context match ({best_score:.2f}) for '{linked}', picking first available."
                        )

                    all_matches = list(direct_pattern.finditer(p.text))

                    if is_generic_bullet:
                        # For generic bullets, rely on reading order (widgets are sorted) and pick the first one
                        m = all_matches[0] if all_matches else None
                    else:
                        m = select_best_match(p.text, all_matches, combined_context)

                    if m:
                        new_text = p.text[: m.start()] + value_to_add.strip() + p.text[m.end() :]
                        apply_replacement(p, new_text)
                        print(
                            f"  > DOCX: Direct bracket replace '{linked}' → '{value_to_add}' in: '{p.text[:80].strip()}...' (score={best_score:.2f}) [Widget Page: {widget.page_number}, Target Page: {p_page}]"
                        )
                        return True

                print(
                    f"  > ❌ DOCX: Could not find direct bracket placeholder '{linked}' (regex: {escaped_linked})"
                )
                return False

            # -------------------------------------------------------------------------
            # 5. Case B — linked_text is an auto-numbered bullet marker e.g. "(1)", "1."
            #    The number is rendered by Word and never appears in para.text
            # -------------------------------------------------------------------------
            is_bullet_like = bool(re.match(r"^[\(\[]?\d+[\)\]\.]?$", linked))
            if is_bullet_like:
                w_num_match = re.search(r"(\d+)", linked)
                if not w_num_match:
                    print(f"  > ❌ DOCX: Could not parse number from bullet anchor '{linked}'")
                    return False

                target_num = int(w_num_match.group(1))

                candidates = []
                for i, p in enumerate(paragraphs):
                    if not has_numPr(p):
                        continue
                    current_num = para_num_map.get(i)
                    if current_num is None or current_num != target_num:
                        continue
                    # has_numPr(p) and current_num == target_num are already guaranteed
                    # by the continue guards above, so no found_match flag is needed.
                    score, details = check_context_similarity(p.text, combined_context, current_num)
                    p_page_num = para_page_map.get(p._p, 1)
                    if enable_page_bonus and max_doc_page > 1:
                        scaled_p_page = p_page_num * page_scale_factor
                        page_diff = abs(scaled_p_page - widget.page_number)
                        if page_diff <= 0.5:
                            score += 2.0
                        elif page_diff <= 1.5:
                            score += 0.5
                    candidates.append((score, p, p_page_num, details))

                candidates.sort(key=lambda x: x[0], reverse=True)

                if candidates:
                    best_score, p, p_page, _ = candidates[0]
                    full_p_text = p.text

                    all_matches = []
                    for pat in placeholder_patterns:
                        m = pat.search(full_p_text, pos=0)
                        if m:
                            all_matches.append(m)

                    match = min(all_matches, key=lambda m: m.start()) if all_matches else None

                    if match:
                        new_text = (
                            full_p_text[: match.start()]
                            + value_to_add.strip()
                            + full_p_text[match.end() :]
                        )
                        apply_replacement(p, new_text)
                    else:
                        new_text = value_to_add.strip() + " " + full_p_text
                        apply_replacement(p, new_text)
                    print(
                        f"  > DOCX: Bullet({target_num}) update → '{value_to_add}' in: '{full_p_text[:80]}' (score={best_score:.2f}) [Widget Page: {widget.page_number}, Target Page: {p_page}]"
                    )
                    return True

                print(f"  > ❌ DOCX: Could not find bullet item number {target_num} for '{linked}'")
                return False

            # -------------------------------------------------------------------------
            # 6. Case C — normal anchor text search e.g. "Founders", "fiscal code"
            # -------------------------------------------------------------------------
            # re.findall with \w+|[^\w\s] never produces whitespace tokens, so the
            # isspace() guard was dead code — removed. Guard against empty linked text
            # to avoid compiling an empty regex that matches everywhere.
            raw_parts = re.findall(r"\w+|[^\w\s]", linked)
            if not raw_parts:
                return False
            flexible_pattern_parts = [re.escape(p) for p in raw_parts]
            anchor_pattern = r"\s*".join(flexible_pattern_parts)
            if re.match(r"^\w", raw_parts[0]):
                anchor_pattern = r"\b" + anchor_pattern
            if re.search(r"\w$", raw_parts[-1]):
                anchor_pattern = anchor_pattern + r"\b"
            anchor_regex = re.compile(anchor_pattern, re.IGNORECASE)

            candidates = []
            for i, p in enumerate(paragraphs):
                full_p_text = p.text
                anchor_match = anchor_regex.search(full_p_text)
                if anchor_match:
                    xml_num = para_num_map.get(i)
                    score, details = check_context_similarity(
                        full_p_text, combined_context, xml_num
                    )

                    # Page alignment bonus
                    p_page_num = para_page_map.get(p._p, 1)
                    if enable_page_bonus and max_doc_page > 1:
                        scaled_p_page = p_page_num * page_scale_factor
                        page_diff = abs(scaled_p_page - widget.page_number)
                        if page_diff <= 0.5:
                            score += 2.0
                        elif page_diff <= 1.5:
                            score += 0.5

                    candidates.append((score, p, anchor_match, p_page_num, details))

            candidates.sort(key=lambda x: x[0], reverse=True)

            if candidates:
                best_score, p, anchor_match, p_page, _ = candidates[0]
                full_p_text = p.text
                search_start = anchor_match.end()

                # Find closest placeholder after anchor
                all_matches = [
                    m
                    for pat in placeholder_patterns
                    for m in [pat.search(full_p_text, pos=search_start)]
                    if m
                ]
                match = min(all_matches, key=lambda m: m.start()) if all_matches else None

                # Discard if placeholder is too far from anchor
                if match and (match.start() - search_start) > 200:
                    match = None

                if match:
                    new_text = (
                        full_p_text[: match.start()]
                        + value_to_add.strip()
                        + full_p_text[match.end() :]
                    )
                    apply_replacement(p, new_text)
                    print(
                        f"  > DOCX: Anchor '{linked}' replaced placeholder → '{value_to_add}' in: '{full_p_text[:80]}' (score={best_score:.2f}) [Widget Page: {widget.page_number}, Target Page: {p_page}]"
                    )
                else:
                    # Anchor found but no nearby placeholder — append directly after anchor
                    new_text = (
                        full_p_text[:search_start] + value_to_add + full_p_text[search_start:]
                    )
                    apply_replacement(p, new_text)
                    print(
                        f"  > DOCX: Anchor '{linked}' appended value → '{value_to_add}' in: '{full_p_text[:80]}' (score={best_score:.2f}) [Widget Page: {widget.page_number}, Target Page: {p_page}]"
                    )
                return True

            # -------------------------------------------------------------------------
            # 7. Case D — Context Fallback: find paragraph with best context match containing a placeholder
            # -------------------------------------------------------------------------
            fallback_candidates = []
            best_debug_candidate = None
            for i, p in enumerate(paragraphs):
                full_p_text = p.text
                # Check if paragraph has any placeholder
                has_placeholder = False
                for pat in placeholder_patterns:
                    if pat.search(full_p_text):
                        has_placeholder = True
                        break

                if not has_placeholder:
                    continue

                xml_num = para_num_map.get(i)
                score, details = check_context_similarity(full_p_text, combined_context, xml_num)

                # Page alignment bonus
                p_page_num = para_page_map.get(p._p, 1)
                if enable_page_bonus and max_doc_page > 1:
                    scaled_p_page = p_page_num * page_scale_factor
                    page_diff = abs(scaled_p_page - widget.page_number)
                    if page_diff <= 0.5:
                        score += 2.0
                    elif page_diff <= 1.5:
                        score += 0.5

                if best_debug_candidate is None or score > best_debug_candidate["score"]:
                    best_debug_candidate = {"score": score, "text": full_p_text, "page": p_page_num}

                # Require a reasonable score to proceed
                if score > 0.6:
                    fallback_candidates.append((score, p, p_page_num, details))

            fallback_candidates.sort(key=lambda x: x[0], reverse=True)

            if fallback_candidates:
                best_score, p, p_page, _ = fallback_candidates[0]
                full_p_text = p.text

                # Find first placeholder in this paragraph
                all_matches = []
                for pat in placeholder_patterns:
                    for m in pat.finditer(full_p_text):
                        all_matches.append(m)

                # Filter overlapping matches by sorting and simple greedy selection
                all_matches.sort(key=lambda m: m.start())
                filtered_matches = []
                last_end = -1
                for m in all_matches:
                    if m.start() >= last_end:
                        filtered_matches.append(m)
                        last_end = m.end()
                all_matches = filtered_matches

                match = select_best_match(full_p_text, all_matches, combined_context)

                if match:
                    # Debug overlap for the paragraph match itself
                    intersection = set(normalize_text(full_p_text).split()).intersection(
                        set(combined_context_norm.split())
                    )

                    new_text = (
                        full_p_text[: match.start()]
                        + value_to_add.strip()
                        + full_p_text[match.end() :]
                    )
                    apply_replacement(p, new_text)
                    print(
                        f"  > DOCX: Context fallback replaced placeholder → '{value_to_add}' in: '{full_p_text[:80].strip()}...' (score={best_score:.2f}) [Widget Page: {widget.page_number}, Target Page: {p_page}]"
                    )
                    print(f"    > Matched context words: {list(intersection)[:10]}")
                    return True

            print(f"  > ❌ DOCX: Could not find anchor '{linked}' for widget")
            if best_debug_candidate:
                intersection = set(
                    normalize_text(best_debug_candidate["text"]).split()
                ).intersection(set(combined_context_norm.split()))
                print(
                    f"    > Best match (score={best_debug_candidate['score']:.2f}, page={best_debug_candidate['page']}) rejected: '{best_debug_candidate['text'][:100].strip()}...' [Widget Page: {widget.page_number}]"
                )
                print(f"    > Match details: {list(intersection)[:10]}")
            return False

        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
                temp_docx.write(source_docx_bytes)
                temp_docx_path = temp_docx.name

            try:
                doc = Document(temp_docx_path)
            except Exception as e:
                # python-docx can raise KeyError/ValueError for missing relationships,
                # zipfile.BadZipFile for invalid archives, or lxml XMLSyntaxError for
                # malformed XML — catch all and attempt LibreOffice repair.
                print(f"Error opening DOCX: {e}. Attempting repair via LibreOffice...")
                try:
                    subprocess.check_call(
                        [
                            "soffice",
                            "--headless",
                            "--convert-to",
                            "docx",
                            temp_docx_path,
                            "--outdir",
                            os.path.dirname(temp_docx_path),
                        ],
                        timeout=60,
                    )
                    doc = Document(temp_docx_path)
                    print("DOCX repaired successfully.")
                except Exception as repair_e:
                    print(f"Repair failed: {repair_e}")
                    raise e

            widgets_added = 0

            # Build page map once
            para_page_map = map_paragraphs_to_pages(doc)
            para_num_map = map_paragraphs_to_list_numbers(doc)
            max_doc_page = max(para_page_map.values()) if para_page_map else 1

            # Calculate scaling factor and determine if page alignment should be used
            page_scale_factor = 1.0
            enable_page_bonus = True
            if total_pdf_pages and max_doc_page > 0:
                if max_doc_page == 1 and total_pdf_pages > 1:
                    print(
                        f"  > DOCX page tags missing (1 vs {total_pdf_pages}). Disabling page alignment."
                    )
                    enable_page_bonus = False
                elif total_pdf_pages != max_doc_page:
                    print(
                        f"  > Page count mismatch: DOCX={max_doc_page}, PDF={total_pdf_pages}. Scaling enabled."
                    )
                    page_scale_factor = total_pdf_pages / max_doc_page

            for w_idx, w in enumerate(widgets):
                # Allow processing even if text_content is empty (to clear placeholders), but require anchor text
                if not w.linked_text:
                    continue

                found_anchor = False
                normalized_anchor = normalize_text(w.linked_text)

                if process_paragraphs(
                    doc.paragraphs,
                    w,
                    para_page_map,
                    max_doc_page,
                    para_num_map,
                    page_scale_factor,
                    enable_page_bonus,
                ):
                    found_anchor = True
                else:
                    for table in doc.tables:
                        if not table.rows:
                            continue

                        tc_grid = self._build_visual_tc_grid(table)

                        if not tc_grid or not tc_grid[0]:
                            continue

                        text_grid = [[self._get_tc_text(tc) for tc in row] for row in tc_grid]

                        header_row_count = 0
                        for r_idx, row_texts in enumerate(text_grid):

                            if not row_texts:
                                continue
                            first_cell_text = normalize_text(row_texts[0])
                            is_identifier = (
                                len(first_cell_text) == 1 and first_cell_text.isalpha()
                            ) or first_cell_text.isdigit()
                            if not is_identifier and len(first_cell_text) > 0:
                                header_row_count = r_idx + 1
                            elif is_identifier:
                                break
                        if header_row_count == 0 and len(text_grid) > 1:
                            header_row_count = 1

                        header_rows = text_grid[:header_row_count]
                        column_headers = [
                            " ".join(
                                filter(
                                    None,
                                    [
                                        header_rows[r][c]
                                        for r in range(header_row_count)
                                        if c < len(header_rows[r])
                                    ],
                                )
                            )
                            for c in range(len(text_grid[0]))
                        ]
                        all_headers_str = " ".join(filter(None, column_headers))

                        anchor_words = set(normalized_anchor.split())
                        if not anchor_words:
                            continue

                        for r_idx, row_texts in enumerate(
                            text_grid[header_row_count:], start=header_row_count
                        ):
                            row_label_text = normalize_text(row_texts[0]) if row_texts else ""
                            row_context_string = f"{all_headers_str} {row_label_text}"
                            # Normalize the entire context string for consistent matching
                            context_words = set(normalize_text(row_context_string).split())

                            # Fuzzy matching: check if at least 75% of anchor words are in context.
                            # This is more robust to OCR errors or slight variations in text.
                            intersection_size = len(anchor_words.intersection(context_words))
                            match_ratio = intersection_size / len(anchor_words)

                            if match_ratio < 0.75:
                                continue
                            print(
                                f"  > Found potential row match for anchor '{normalized_anchor}' in table row {r_idx}"
                            )
                            best_match_score, target_col_idx = -1, -1

                            for c_idx, h_text in enumerate(column_headers):
                                if not h_text:
                                    continue
                                score = len(anchor_words.intersection(set(h_text.split())))
                                if score > best_match_score:
                                    best_match_score = score
                                    target_col_idx = c_idx

                            if target_col_idx != -1:
                                value_to_add = ""
                                if w.label in ["checkbox", "radio"]:
                                    value_to_add = " \u2611" if w.is_filled else " \u2610"
                                elif w.label == "signature" and w.text_content:
                                    value_to_add = f" {w.text_content}"
                                elif w.text_content:
                                    value_to_add = f" {w.text_content.strip()}"
                                target_tc = tc_grid[r_idx][target_col_idx]
                                if target_tc is not None:
                                    self._set_tc_text(target_tc, value_to_add)
                                    print(
                                        f"  > Successfully filled cell at ({r_idx}, {target_col_idx}) for anchor '{w.linked_text}'"
                                    )
                                    found_anchor = True
                                    break
                        if found_anchor:
                            break

                if found_anchor:
                    widgets_added += 1
                else:
                    msg = f"Could not find anchor '{w.linked_text}' (normalized: '{normalized_anchor}') for widget {w_idx}"
                    print(f"  > ❌ DOCX: {msg}")
                    warnings.append(msg)

            print(f"Added {widgets_added} values to the DOCX.")
            doc.save(output_docx_path)
            print(f"Augmented DOCX saved to {output_docx_path}")

        except ImportError:
            msg = "Error augmenting DOCX: `python-docx` library not found. Please install it."
            print(msg)
            warnings.append(msg)
        except Exception as e:
            msg = f"Error augmenting DOCX: {e}"
            print(msg)
            import traceback

            traceback.print_exc()
            warnings.append(msg)
        finally:
            if temp_docx_path and os.path.exists(temp_docx_path):
                os.remove(temp_docx_path)
        return warnings


FORM_FILLING_PROMPTS = {
    "detection": {
        "system": "You are an expert form parser. Your task is to analyze a document page image and identify all form widgets present. Images dimensions should be mapped to 1000x1000 coordinate space, with (0,0) at the top-left. ",
        "user": "Inspect the provided image of a document page and identify and enumerate all form widgets present, use visual cues to identify all the areas where data could be input or has been entered. "
        "Ground the bounding boxes of the detected widgets to the image using the provided coordinates. "
        "The detected widgets will be overlayed over the form image. "
        "Return a JSON object with a key 'widgets' containing a list of detected widgets.\n"
        "Each item must have:\n"
        "'label' (one of: 'text_input', 'checkbox', 'radio', 'signature'),\n"
        "'box_2d' (bounding box in [ymin, xmin, ymax, xmax] format, normalized 0-1000 for both dimensions),\n"
        "'is_filled' (boolean),\n"
        "'text_content' (string, the text value inside the widget if present),\n"
        "'description' (short text describing the widget context).\n\n"
        "Identify if the widget has a label (e.g. agency, code, ...) and, revise the box_2d to ensure that the label is not included in the bounding box, to prevent that it is not visible when filling the form, but none of the detected widgets should be discarded because of this requirements, it should be possible to have widgets without a label nearby.",
    },
    "filling": {
        "system": (
            "You are an AI assistant that fills forms based on provided information. "
            "You will receive a list of fields and a text containing data. "
            "Map the data to the fields. Use semantic reasoning to match data to field labels even if they are not identical. "
            "For 'text_input', provide the string value. "
            "For 'checkbox' or 'radio', provide a boolean (true for checked/yes, false for unchecked/no). "
            "Return a JSON object where keys are the field 'id' and values are the assigned values."
        ),
        "user": "Data:\n{prompt_text}\n\nFields:\n{fields}",
    },
    "analyzer": {
        "system": (
            "You are a form analyzer. I will provide an image of a form where the widgets to analyze are highlighted with red bounding boxes and labeled with their IDs in red. "
            "I will also provide a list of these widgets with their IDs and bounding box coordinates (normalized 0-1000). "
            "The form image dimensions should be within a 1000x1000 coordinate space, so the bounding box coordinates correspond to positions on the image. "
            "You will ground the bounding boxes to the form image using both the visual red boxes/IDs and the coordinates. "
            "For each field identified by its ID, identify the semantic label (e.g., the text 'Name:' next to the box) "
            "and provide a short description of what should be entered. "
            "Also extract the surrounding text context (sentences, headers, or paragraphs containing this field) to help locate it in the document text. "
            "Use visual cues to target the specific label for each widget. "
            "Do not use text that belongs to other nearby widgets. "
            "Return a JSON object with a key 'fields' containing a list of objects: {'id': int, 'label': string, 'description': string, 'surrounding_text': string}."
        ),
        "user": "Widgets to analyze:\n{widgets}",
    },
}

FORM_FILLING_SCHEMAS = {
    "detection": {
        "type": "object",
        "properties": {
            "widgets": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "label": {
                            "type": "string",
                            "enum": ["text_input", "checkbox", "radio", "signature"],
                        },
                        "box_2d": {
                            "type": "array",
                            "items": {"type": "integer", "minimum": 0, "maximum": 1000},
                            "minItems": 4,
                            "maxItems": 4,
                        },
                        "is_filled": {"type": "boolean"},
                        "text_content": {"type": "string"},
                        "description": {"type": "string"},
                    },
                    "required": ["label", "box_2d", "is_filled"],
                },
            }
        },
        "required": ["widgets"],
    },
    "analyzer": {
        "type": "object",
        "properties": {
            "fields": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "id": {"type": "integer"},
                        "label": {"type": "string"},
                        "description": {"type": "string"},
                        "surrounding_text": {"type": "string"},
                    },
                    "required": ["id", "label", "description"],
                },
            }
        },
        "required": ["fields"],
    },
}


def get_form_filling_prompts(task: str) -> tuple[str, str]:
    try:
        prompt_data = FORM_FILLING_PROMPTS[task]
    except KeyError:
        raise ValueError(f"Invalid form filling task: {task}")
    return prompt_data["system"], prompt_data["user"]


def get_form_filling_schema(task: str) -> Dict:
    try:
        return FORM_FILLING_SCHEMAS[task]
    except KeyError:
        raise ValueError(f"Invalid form filling schema task: {task}")


class FormDetector:
    """
    Wrapper for Gemini and local models to detect form widgets in images.
    """

    def __init__(
        self,
        gemini_model_name: str = "gemini-3-flash-preview",
    ):
        self.gemini_model_name = gemini_model_name

    def detect_gemini(
        self,
        page_image_path: str,
        page_width: float,
        page_height: float,
    ) -> tuple[list[DetectedWidget], int, int]:
        """
        Runs Gemini detection.
        """
        from google.genai import types
        from tensorlake_docai.providers.model_provider_utils import _make_gemini_call

        input_tokens = 0
        output_tokens = 0
        if not os.path.exists(page_image_path):
            print(
                f"Skipping Gemini detection (Path: {page_image_path})",
            )
            return [], 0, 0

        print(f"Running Gemini detection on {page_image_path}...")

        system_prompt, user_prompt = get_form_filling_prompts("detection")

        detection_schema = get_form_filling_schema("detection")

        try:
            with Image.open(page_image_path) as img:
                config_overrides = {
                    "thinking_config": types.ThinkingConfig(thinking_level="high"),
                    "temperature": 0.0,
                }

                response_text, in_tok, out_tok = asyncio.run(
                    _make_gemini_call(
                        system_instruction=system_prompt,
                        user_prompt=user_prompt,
                        timeout=360,
                        images=[img],
                        model_name=self.gemini_model_name,
                        job_type="structured_extraction",
                        config_overrides=config_overrides,
                        json_schema=json.dumps(detection_schema),
                    )
                )
                input_tokens += in_tok
                output_tokens += out_tok
            data = json.loads(response_text)

            if isinstance(data, list):
                data = data[0] if data else {}

            widgets = []
            for w in data.get("widgets", []):
                box = w["box_2d"]
                # Convert [ymin, xmin, ymax, xmax] to [x1, y1, x2, y2] in page points.
                y1 = (box[0] / 1000) * page_height
                x1 = (box[1] / 1000) * page_width
                y2 = (box[2] / 1000) * page_height
                x2 = (box[3] / 1000) * page_width

                label = w.get("label", "text_input")

                # Heuristic: Expand checkboxes slightly as they are often detected too tightly
                if label in ["checkbox", "radio"]:
                    w_box = x2 - x1
                    x1 -= w_box * 0.1
                    x2 += w_box * 0.1

                text_content = w.get("text_content")
                if isinstance(text_content, str) and text_content.lower() == "null":
                    text_content = ""

                widgets.append(
                    DetectedWidget(
                        label=label,
                        score=1.0,
                        bbox=BoundingBox(x1, y1, x2, y2),
                        is_filled=w.get("is_filled", False),
                        description=w.get("description"),
                        text_content=text_content or "",
                    ),
                )

            return widgets, input_tokens, output_tokens
        except Exception as e:
            print(f"Gemini detection failed: {e}")
            return [], 0, 0


class FormFiller:
    """
    Uses an LLM to map a user prompt (data) to the detected form widgets.
    """

    def __init__(self, model_name: str = "gemini-3-flash-preview"):
        self.model_name = model_name

    def fill(self, widgets: list[DetectedWidget], prompt_text: str) -> tuple[int, int]:
        """
        Fills the widgets based on the provided prompt text.
        Updates the 'text_content' or 'is_filled' attributes of the widgets.
        """
        from tensorlake_docai.providers.model_provider_utils import _make_gemini_call

        input_tokens = 0
        output_tokens = 0
        if not widgets:
            return 0, 0

        properties = {}
        fields_desc = []
        for w in widgets:
            fields_desc.append(
                {
                    "id": w.field_name,
                    "label": w.linked_text or w.label,
                    "type": w.label,
                    "description": w.description,
                }
            )

            if w.label in ["checkbox", "radio"]:
                prop_type = "boolean"
            else:
                prop_type = "string"

            if w.field_name:
                properties[w.field_name] = {"type": prop_type}

        system_prompt, user_prompt = get_form_filling_prompts("filling")

        user_prompt = user_prompt.format(
            prompt_text=prompt_text, fields=json.dumps(fields_desc, indent=2)
        )

        try:
            response_text, in_tok, out_tok = asyncio.run(
                _make_gemini_call(
                    user_prompt=user_prompt,
                    images=[],
                    model_name=self.model_name,
                    system_instruction=system_prompt,
                    job_type="structured_extraction",
                    config_overrides={"temperature": 0.0},
                    # json_schema=json.dumps(filling_schema),
                )
            )
            input_tokens += in_tok
            output_tokens += out_tok

            result = json.loads(response_text)
            print(f"  LLM Response for filling: {json.dumps(result, indent=2)}")

            for w in widgets:
                if w.field_name in result:
                    val = result[w.field_name]
                    print(f"Filling field '{w.field_name}' with value: {val}")
                    if w.label in ["checkbox", "radio"]:
                        if isinstance(val, bool):
                            w.is_filled = val
                        elif isinstance(val, str):
                            w.is_filled = val.lower().strip() in [
                                "true",
                                "yes",
                                "checked",
                                "x",
                                "on",
                                "selected",
                            ]
                    else:
                        if val is not None:
                            str_val = str(val)
                            if str_val.lower() == "null":
                                w.text_content = ""
                            else:
                                w.text_content = str_val

        except Exception as e:
            print(f"Error during form filling: {e}")
        return input_tokens, output_tokens


class MetadataRefiner:
    """
    Uses a VLM to analyze widgets that lack proper labels (e.g. existing PDF fields
    with generic names) and assigns them semantic labels based on visual context.
    """

    def __init__(self, model_name: str = "gemini-3-flash-preview"):
        self.model_name = model_name

    def refine(
        self,
        widgets: list[DetectedWidget],
        image_path: str,
        width: float,
        height: float,
        extract_context: bool = False,
    ) -> tuple[int, int]:
        from tensorlake_docai.providers.model_provider_utils import _make_gemini_call
        from google.genai import types

        input_tokens = 0
        output_tokens = 0
        if not widgets:
            return 0, 0

        if not image_path or not os.path.exists(image_path):
            print(
                f"  Warning: Skipping metadata refinement for {len(widgets)} widgets due to missing image."
            )
            return 0, 0

        # Identify widgets that need refinement (missing linked_text or generic names)
        widgets_to_refine = []
        for i, w in enumerate(widgets):
            # Refine if no linked text, or if the field name looks like a generic ID
            if (
                extract_context
                or not w.linked_text
                or (w.field_name and "Check Box" in w.field_name)
            ):
                widgets_to_refine.append((i, w))

        if not widgets_to_refine:
            return 0, 0

        print(f"  Refining metadata for {len(widgets_to_refine)} widgets using VLM...")

        for i, w in widgets_to_refine:
            print(
                f"    - Queuing widget for refinement (id: {i}, name: '{w.field_name}', label: '{w.linked_text}')"
            )

        widget_inputs = []
        for idx, w in widgets_to_refine:
            # Normalize coordinates to 0-1000
            ymin = int((w.bbox.y1 / height) * 1000)
            xmin = int((w.bbox.x1 / width) * 1000)
            ymax = int((w.bbox.y2 / height) * 1000)
            xmax = int((w.bbox.x2 / width) * 1000)
            widget_inputs.append({"id": idx, "box_2d": [ymin, xmin, ymax, xmax]})

        analyzer_schema = get_form_filling_schema("analyzer")

        system_prompt, user_prompt = get_form_filling_prompts("analyzer")

        try:
            with Image.open(image_path) as img:
                # Draw bounding boxes and IDs on the image
                from PIL import ImageDraw, ImageFont

                draw_img = img.copy()
                draw = ImageDraw.Draw(draw_img)

                try:
                    try:
                        font = ImageFont.truetype("DejaVuSans.ttf", 20)
                    except OSError:
                        try:
                            font = ImageFont.load_default(size=20)
                        except TypeError:
                            font = ImageFont.load_default()
                except Exception:
                    font = None

                img_w, img_h = draw_img.size
                scale_x = img_w / width if width > 0 else 1.0
                scale_y = img_h / height if height > 0 else 1.0

                for idx, w in widgets_to_refine:
                    x1 = w.bbox.x1 * scale_x
                    y1 = w.bbox.y1 * scale_y
                    x2 = w.bbox.x2 * scale_x
                    y2 = w.bbox.y2 * scale_y

                    draw.rectangle([x1, y1, x2, y2], outline="red", width=3)

                    text = str(idx)
                    text_x, text_y = x1, y1
                    if font:
                        try:
                            left, top, right, bottom = draw.textbbox((0, 0), text, font=font)
                            text_w = right - left
                            text_h = bottom - top
                            text_x = x1 + (x2 - x1 - text_w) / 2
                            text_y = y1 + (y2 - y1 - text_h) / 2
                        except AttributeError:
                            pass
                        draw.text((text_x, text_y), text, fill="red", font=font)
                    else:
                        draw.text((x1, y1), text, fill="red")

                config_overrides = {
                    "thinking_config": types.ThinkingConfig(thinking_level="high"),
                    "temperature": 0.0,
                }

                user_prompt = user_prompt.format(widgets=json.dumps(widget_inputs))
                response_text, in_tok, out_tok = asyncio.run(
                    _make_gemini_call(
                        system_instruction=system_prompt,
                        user_prompt=user_prompt,
                        images=[draw_img],
                        model_name=self.model_name,
                        job_type="structured_extraction",
                        config_overrides=config_overrides,
                        json_schema=json.dumps(analyzer_schema),
                    )
                )
                input_tokens += in_tok
                output_tokens += out_tok
            result = json.loads(response_text)

            for item in result.get("fields", []):
                idx = item.get("id")
                if idx is not None and 0 <= idx < len(widgets):
                    w = widgets[idx]
                    if item.get("label") and item["label"] != w.linked_text:
                        print(
                            f"    - Refining widget {idx}: setting label to '{item['label']}' (was '{w.linked_text}')"
                        )
                        w.linked_text = item["label"]
                    if item.get("description") and item["description"] != w.description:
                        print(
                            f"    - Refining widget {idx}: setting description to '{item['description']}' (was '{w.description}')"
                        )
                        w.description = item["description"]
                    if item.get("surrounding_text"):
                        w.surrounding_text = item["surrounding_text"]

        except Exception as e:
            print(f"  Metadata refinement failed: {e}")
        return input_tokens, output_tokens


class FormFillingSystem:
    def __init__(self):
        self.detector = FormDetector()
        self.associator = LabelAssociator()
        self.augmenter = PdfFormAugmenter()
        self.filler = FormFiller()
        self.refiner = MetadataRefiner()
        self.docx_augmenter = DocxFormAugmenter()

    def get_page_mediabox(self, pdf_path: str, page_number: int):
        """Gets page mediabox and rotation from PDF in a single read."""
        import pypdf

        try:
            reader = pypdf.PdfReader(pdf_path)
            if 1 <= page_number <= len(reader.pages):
                page = reader.pages[page_number - 1]

                rotation = 0
                if hasattr(page, "rotation") and page.rotation is not None:
                    rotation = page.rotation
                elif "/Rotate" in page:
                    rotation = page["/Rotate"]
                rotation = int(rotation) % 360

                # Prefer CropBox if available, as it matches the visible area rendered by pdf2image
                mediabox = page.cropbox if "/CropBox" in page else page.mediabox
                return mediabox, rotation
        except Exception as e:
            print(f"Error reading PDF mediabox: {e}")
        return None, 0

    def get_existing_widgets(
        self,
        pdf_path: str,
        page_number: int,
        page_mediabox,
        rotation: int = 0,
        expected_width: float = None,
        expected_height: float = None,
    ) -> list[DetectedWidget]:
        """Extracts existing widgets from PDF annotations."""
        import pypdf

        widgets = []
        if not page_mediabox:
            return []

        page_left = float(page_mediabox.left)
        page_bottom = float(page_mediabox.bottom)
        page_width = float(page_mediabox.width)
        page_height = float(page_mediabox.height)

        # Determine current dimensions based on rotation
        if rotation in [90, 270]:
            current_width = page_height
            current_height = page_width
        else:
            current_width = page_width
            current_height = page_height

        scale_x = expected_width / current_width if expected_width and current_width > 0 else 1.0
        scale_y = (
            expected_height / current_height if expected_height and current_height > 0 else 1.0
        )

        try:
            reader = pypdf.PdfReader(pdf_path)
            if 1 <= page_number <= len(reader.pages):
                page = reader.pages[page_number - 1]
                if "/Annots" in page:
                    for annot in page["/Annots"]:
                        obj = annot.get_object()
                        if obj.get("/Subtype") == "/Widget":
                            field_name = obj.get("/T")
                            rect = obj.get("/Rect")  # [x_ll, y_ll, x_ur, y_ur]

                            # Convert PDF (Bottom-Left) to Image (Top-Left) coords
                            x_ll, y_ll, x_ur, y_ur = [float(c) for c in rect]

                            # Normalize to 0-based coordinates relative to mediabox
                            x_min = x_ll - page_left
                            y_min = y_ll - page_bottom
                            x_max = x_ur - page_left
                            y_max = y_ur - page_bottom

                            # Apply rotation transform
                            if rotation == 0:
                                x1 = x_min
                                y1 = page_height - y_max
                                x2 = x_max
                                y2 = page_height - y_min
                            elif rotation == 90:
                                x1 = y_min
                                y1 = x_min
                                x2 = y_max
                                y2 = x_max
                            elif rotation == 180:
                                x1 = page_width - x_max
                                y1 = y_min
                                x2 = page_width - x_min
                                y2 = y_max
                            elif rotation == 270:
                                x1 = page_height - y_max
                                y1 = page_width - x_max
                                x2 = page_height - y_min
                                y2 = page_width - x_min
                            else:
                                x1 = x_min
                                y1 = page_height - y_max
                                x2 = x_max
                                y2 = page_height - y_min

                            # Normalize bbox (x1 < x2, y1 < y2)
                            x1, x2 = min(x1, x2), max(x1, x2)
                            y1, y2 = min(y1, y2), max(y1, y2)

                            # Scale to expected dimensions
                            x1 *= scale_x
                            x2 *= scale_x
                            y1 *= scale_y
                            y2 *= scale_y

                            # Determine type
                            label = "text_input"
                            ft = obj.get("/FT")
                            if ft == "/Btn":
                                # Check flags for radio vs checkbox if needed, simplified here
                                label = "checkbox"
                            elif ft == "/Sig":
                                label = "signature"

                            # Extract value and filled status from PDF data
                            text_content = None
                            is_filled = False
                            if "/V" in obj:
                                raw_val = obj.get_object().get("/V")
                                if isinstance(raw_val, pypdf.generic.TextStringObject):
                                    text_content = str(raw_val)
                                    if text_content:
                                        is_filled = True
                                elif isinstance(raw_val, pypdf.generic.NameObject):
                                    if raw_val not in ("/Off", ""):
                                        is_filled = True
                                    if raw_val == "/Yes":
                                        text_content = "checked"
                                    elif raw_val != "/Off":
                                        # Clean up the name object string
                                        text_content = str(raw_val).lstrip("/")

                            widgets.append(
                                DetectedWidget(
                                    label=label,
                                    score=1.0,
                                    bbox=BoundingBox(x1, y1, x2, y2),
                                    is_filled=is_filled,
                                    text_content=text_content or "",
                                    is_existing=True,
                                    field_name=str(field_name) if field_name else None,
                                ),
                            )
        except Exception as e:
            print(f"Error extracting existing widgets: {e}")

        return widgets

    def process_page(
        self,
        page_data: PageData,
        source_pdf: str = "document.pdf",
        use_acroform: bool = True,
        use_widget_detection: bool = True,
        is_docx: bool = False,
    ) -> tuple[list[DetectedWidget], int, int]:
        print(f"Processing Page {page_data.page_number}...")
        total_input_tokens = 0
        total_output_tokens = 0

        # 0. Get Page Dimensions (mediabox and rotation read in a single PDF open)
        mediabox, rotation = self.get_page_mediabox(source_pdf, page_data.page_number)
        if mediabox:
            w, h = float(mediabox.width), float(mediabox.height)
            if rotation in [90, 270]:
                w, h = h, w
        else:
            # Fallback if PDF read fails, assume standard letter
            w, h = 612, 792

        # 1. Get Existing Widgets (Precedence)
        final_widgets = []
        if use_acroform:
            existing_widgets = self.get_existing_widgets(
                source_pdf,
                page_data.page_number,
                mediabox,
                rotation=rotation,
                expected_width=w,
                expected_height=h,
            )
            final_widgets.extend(existing_widgets)
            print(f"  Found {len(existing_widgets)} existing widgets in PDF.")

        # Helper to check for duplicates against accepted widgets
        def is_duplicate(widget, accepted_widgets):
            for accepted in accepted_widgets:
                # Check IoU
                if widget.bbox.iou(accepted.bbox) > 0.1:
                    return True
                # Check Containment (one inside the other)
                if widget.bbox.is_contained_in(accepted.bbox) or accepted.bbox.is_contained_in(
                    widget.bbox
                ):
                    return True
            return False

        # 3. Widget Detection
        if use_widget_detection and page_data.image_path and os.path.exists(page_data.image_path):
            gemini_widgets, in_tok, out_tok = self.detector.detect_gemini(
                page_data.image_path, w, h
            )
            total_input_tokens += in_tok
            total_output_tokens += out_tok
            added_count = 0
            for gw in gemini_widgets:
                if not is_duplicate(gw, final_widgets):
                    final_widgets.append(gw)
                    added_count += 1
            print(
                f"  Added {added_count} widgets from Widget Detection (discarded {len(gemini_widgets) - added_count} duplicates)."
            )
        elif use_widget_detection and page_data.image_path:
            print(f"  Image path specified but not found: {page_data.image_path}")

        print(f"  Total widgets for page: {len(final_widgets)}")

        # 3.5 Refine Metadata (VLM)
        # Use VLM to fix missing labels or generic names (e.g. "Check Box3")
        # If is_docx is True, we force refinement on all widgets to get better context
        in_tok, out_tok = self.refiner.refine(
            final_widgets, page_data.image_path, w, h, extract_context=is_docx
        )
        total_input_tokens += in_tok
        total_output_tokens += out_tok

        # 4. Associate Labels
        for widget in final_widgets:
            widget.page_number = page_data.page_number

            # Extract surrounding text from fragments (geometric)
            geo_context = self._extract_geometric_context(widget, page_data.fragments)

            # Combine VLM context (if any) with geometric context
            if widget.surrounding_text:
                # Prioritize VLM context but keep geometric as fallback/augmentation
                widget.surrounding_text = f"{widget.surrounding_text} {geo_context}"
            else:
                widget.surrounding_text = geo_context

            # Only associate if a label hasn't already been assigned by the refiner
            if not widget.linked_text:
                linked_text = self.associator.associate(
                    widget, page_data.fragments, is_docx=is_docx
                )
                widget.linked_text = linked_text

        # Sort widgets by position (Top-Left to Bottom-Right) to ensure reading order matching
        # This is critical for matching multiple identical placeholders in order
        final_widgets.sort(key=lambda w: (int(w.bbox.y1 / 10), int(w.bbox.x1)))

        return final_widgets, total_input_tokens, total_output_tokens

    def assign_field_names(self, widgets: list[DetectedWidget]):
        """Assigns deterministic field names to widgets before processing."""
        for w in widgets:
            if w.field_name:
                continue

            unique_suffix = f"{w.page_number}_{int(w.bbox.x1)}_{int(w.bbox.y1)}"
            if w.linked_text:
                safe_label = "".join(c for c in w.linked_text if c.isalnum() or c in ("_", "-"))
                field_name = f"{safe_label}_{unique_suffix}"
            else:
                field_name = f"field_{unique_suffix}"
            w.field_name = field_name

    def _extract_geometric_context(
        self, widget: DetectedWidget, fragments: list[PageFragment]
    ) -> str:
        """Finds text spatially close to the widget to use as context for anchoring."""
        w_center_y = widget.bbox.center[1]
        w_height = widget.bbox.height

        line_fragments = []
        for frag in fragments:
            if frag.fragment_type not in ["text", "title", "section_header", "list_item"]:
                continue

            f_center_y = frag.bbox.center[1]
            # Check if vertically aligned (roughly same line)
            if abs(f_center_y - w_center_y) < (max(w_height, frag.bbox.height) * 1.5):
                line_fragments.append(frag)

        # Sort by x coordinate to reconstruct the line
        line_fragments.sort(key=lambda f: f.bbox.x1)
        return " ".join([f.content for f in line_fragments])


# --- Ingestion Helpers ---


SECRETS = [
    "AWS_ACCESS_KEY_ID",
    "AWS_SECRET_ACCESS_KEY",
    "AWS_REGION",
]


@cls()
class FormFilling:
    @function(
        image=file_convertion_image,
        description="Fill a PDF form using AI.",
        timeout=30 * 60,  # 30 minutes
        cpu=2,
        memory=8,
        # output_encoder = "json"
        # The function is not using /tmp disk space, just reserve a small amount
        ephemeral_disk=2,
        secrets=SECRETS,
        min_containers=int(os.getenv("TENSORLAKE_MIN_CONTAINERS", "0")),
    )
    def run(self, result: ParseResult) -> ParsedDocumentRef:
        import pypdf

        request = result.request.form_filling

        if result.docx_converted_pdf_base64:
            source_pdf = base64.b64decode(result.docx_converted_pdf_base64)
        else:
            source_pdf = result.request.file_bytes
            if isinstance(source_pdf, str):
                try:
                    source_pdf = base64.b64decode(source_pdf)
                except Exception:
                    pass

        if not source_pdf:
            raise ValueError("source_pdf is required")

        fill_prompt = request.fill_prompt
        ignore_source_values = request.ignore_source_values
        no_acroform = request.no_acroform
        no_widget_detection = request.no_widget_detection

        is_docx_source = bool(result.docx_converted_pdf_base64)

        total_input_tokens = 0
        total_output_tokens = 0

        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = os.path.join(temp_dir, "input.pdf")
            output_pdf_path = os.path.join(temp_dir, "output.pdf")

            with open(input_path, "wb") as f:
                f.write(source_pdf)

            system = FormFillingSystem()

            from pdf2image import convert_from_path

            reader = pypdf.PdfReader(input_path)
            total_source_pdf_pages = len(reader.pages)
            pages_to_parse = (
                list(range(1, total_source_pdf_pages + 1))
                if not result.request.pages_to_parse
                else result.request.pages_to_parse
            )

            # Create a map for easy lookup
            pages_data_map = {
                i: PageData(page_number=i, fragments=[], image_path=None) for i in pages_to_parse
            }

            # Populate fragments from document_layout
            if result.document_layout:
                for page_layout in result.document_layout.pages:
                    if page_layout.page_number in pages_data_map:
                        page_data = pages_data_map[page_layout.page_number]
                        for element in page_layout.elements:
                            if element.ocr_text:
                                bbox_tuple = element.bbox
                                page_data.fragments.append(
                                    PageFragment(
                                        fragment_type=element.fragment_type.value,
                                        content=element.ocr_text,
                                        bbox=BoundingBox(
                                            x1=bbox_tuple[0],
                                            y1=bbox_tuple[1],
                                            x2=bbox_tuple[2],
                                            y2=bbox_tuple[3],
                                        ),
                                        reading_order=element.reading_order,
                                    )
                                )

            pages = list(pages_data_map.values())

            try:
                images = convert_from_path(input_path)
                for page in pages:
                    img_idx = page.page_number - 1
                    if 0 <= img_idx < len(images):
                        image_path = os.path.join(temp_dir, f"page_{page.page_number}.jpg")
                        images[img_idx].save(image_path, "JPEG")
                        page.image_path = image_path
            except Exception as e:
                print(f"Warning: PDF rendering failed: {e}")

            all_widgets = []
            pages_processed_count = 0
            for page in pages:
                pages_processed_count += 1
                widgets, in_tok, out_tok = system.process_page(
                    page,
                    source_pdf=input_path,
                    use_acroform=not no_acroform,
                    use_widget_detection=not no_widget_detection,
                    is_docx=is_docx_source,
                )
                total_input_tokens += in_tok
                total_output_tokens += out_tok
                all_widgets.extend(widgets)

            if ignore_source_values:
                for w in all_widgets:
                    if w.is_existing:
                        w.text_content = None
                        w.is_filled = False

            system.assign_field_names(all_widgets)

            if fill_prompt:
                in_tok, out_tok = system.filler.fill(all_widgets, fill_prompt)
                total_input_tokens += in_tok
                total_output_tokens += out_tok

            system.augmenter.augment_pdf(input_path, output_pdf_path, all_widgets)

            out_pdf_b64 = None
            if os.path.exists(output_pdf_path):
                with open(output_pdf_path, "rb") as f:
                    out_pdf_bytes = f.read()
                out_pdf_b64 = base64.b64encode(out_pdf_bytes).decode("utf-8")

            filled_docx_base64 = None
            docx_warnings = []

            if result.docx_converted_pdf_base64:
                # If input was DOCX, perform native DOCX augmentation
                print("Performing native DOCX augmentation...")
                output_docx_path = os.path.join(temp_dir, "output.docx")

                # Use original DOCX bytes from request
                docx_input_bytes = result.request.file_bytes
                if isinstance(docx_input_bytes, str):
                    try:
                        docx_input_bytes = base64.b64decode(docx_input_bytes)
                    except Exception as e:
                        print(f"Error decoding DOCX bytes: {e}")
                        docx_input_bytes = None
                if docx_input_bytes:
                    print(f"DOCX input bytes length: {len(docx_input_bytes)}")
                    docx_warnings = system.docx_augmenter.augment_docx(
                        docx_input_bytes,
                        output_docx_path,
                        all_widgets,
                        total_pdf_pages=total_source_pdf_pages,
                    )
                    if os.path.exists(output_docx_path):
                        with open(output_docx_path, "rb") as f:
                            docx_bytes = f.read()
                        filled_docx_base64 = base64.b64encode(docx_bytes).decode("utf-8")
                        print(f"Filled DOCX base64 generated, length: {len(filled_docx_base64)}")
                    else:
                        print("Output DOCX file was not created.")
                        docx_warnings.append("Output DOCX file was not created.")

                else:
                    print("No DOCX input bytes found in request.")
                    docx_warnings.append("No DOCX input bytes found in request.")

            widgets_data = [asdict(w) for w in all_widgets]
            metadata = {"detected_widgets": widgets_data, "docx_fill_warnings": docx_warnings}

            usage = Usage(
                pages_parsed=pages_processed_count,
                extraction_input_tokens_used=total_input_tokens,
                extraction_output_tokens_used=total_output_tokens,
            )

            result.usage = usage
            result.form_filling_result = FormFillingResult(
                filled_pdf_base64=out_pdf_b64,
                filled_docx_base64=filled_docx_base64,
                metadata=metadata,
            )

            return format_final_output(result)
